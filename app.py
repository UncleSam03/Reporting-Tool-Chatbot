import os
import json
from flask import Flask, request, jsonify, send_file, render_template_string
from docx import Document

app = Flask(__name__)

# Constants
TEMPLATE_PATH = r"c:\Users\samuk\Desktop\Reporting Tool Chatbot\HFF SUPPORT GROUP REPORTING TOOL.docx"
OUTPUT_PATH = r"c:\Users\samuk\Desktop\Reporting Tool Chatbot\HFF SUPPORT GROUP REPORT FORM_Filled.docx"

# Conversational flow definition
FLOW = [
    {
        "id": "facilitator",
        "question": "Hello! Welcome to the HFF Support Group Reporting Chatbot. Let's start compiling your monthly report.\n\nWhat is your name (Group Facilitator)?",
        "type": "text",
        "label": "Facilitator Name"
    },
    {
        "id": "cell_num",
        "question": "Thank you. What is your Cell # (Phone number)?",
        "type": "text",
        "label": "Cell Number"
    },
    {
        "id": "town_village",
        "question": "Got it. Which Town or Village is this support group located in?",
        "type": "text",
        "label": "Town/Village"
    },
    {
        "id": "location",
        "question": "What is the specific Location or Venue of the meetings?",
        "type": "text",
        "label": "Specific Location"
    },
    {
        "id": "meeting_day",
        "question": "On which day of the week does the group meet?\n\n1. Monday\n2. Tuesday\n3. Wednesday\n4. Thursday\n5. Friday\n6. Saturday\n7. Sunday",
        "type": "choice",
        "choices": ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"],
        "label": "Meeting Day"
    },
    {
        "id": "meeting_time",
        "question": "At what time do meetings start? (e.g. 10:00 AM or 14:00)",
        "type": "text",
        "label": "Meeting Time"
    },
    {
        "id": "month",
        "question": "Which Month is this report for?\n\n1. January\n2. February\n3. March\n4. April\n5. May\n6. June\n7. July\n8. August\n9. September\n10. October\n11. November\n12. December",
        "type": "choice",
        "choices": ["January", "February", "March", "April", "May", "June", "July", "August", "September", "October", "November", "December"],
        "label": "Reporting Month"
    },
    {
        "id": "met_status",
        "question": "Were you able to meet this month?\n\n1. Yes\n2. No",
        "type": "choice",
        "choices": ["Yes", "No"],
        "label": "Able to Meet?"
    },
    # Conditional fields (Yes)
    {
        "id": "attendees_male",
        "question": "How many Male participants turned up?",
        "type": "integer",
        "condition": lambda ans: ans.get("met_status") == "Yes",
        "label": "Male Attendees"
    },
    {
        "id": "attendees_female",
        "question": "How many Female participants turned up?",
        "type": "integer",
        "condition": lambda ans: ans.get("met_status") == "Yes",
        "label": "Female Attendees"
    },
    {
        "id": "lessons_interesting",
        "question": "How do you make the lessons interesting for the participants?",
        "type": "text",
        "condition": lambda ans: ans.get("met_status") == "Yes",
        "label": "Lessons Interest"
    },
    # General fields again
    {
        "id": "challenges",
        "question": "What challenges have you encountered? (Select all that apply, separated by commas, e.g. 1,3,5 or type 'None'):\n\n1. Attendance / Punctuality issues\n2. Lack of learning materials / manuals\n3. Venue availability or bad weather\n4. Disinterest / low participation from members\n5. Transportation constraints\n6. Healthcare or physical health barriers\n7. Other",
        "type": "multi_choice",
        "choices": [
            "Attendance / Punctuality issues",
            "Lack of learning materials / manuals",
            "Venue availability or bad weather",
            "Disinterest / low participation from members",
            "Transportation constraints",
            "Healthcare or physical health barriers",
            "Other"
        ],
        "label": "Challenges Encountered"
    },
    {
        "id": "challenges_other",
        "question": "You selected 'Other'. Please describe the other challenges you encountered:",
        "type": "text",
        "condition": lambda ans: "Other" in ans.get("challenges", []),
        "label": "Other Challenges Description"
    },
    {
        "id": "challenges_resolved",
        "question": "How were the challenges resolved?",
        "type": "text",
        "label": "How Challenges Resolved"
    },
    {
        "id": "challenges_unresolved",
        "question": "If the challenges haven't been resolved, how do you think they can be resolved?",
        "type": "text",
        "label": "Unresolved Resolution Ideas"
    },
    {
        "id": "add_testimony",
        "question": "Would you like to add a Support Group Participant Testimony to this report?\n\n1. Yes\n2. No",
        "type": "choice",
        "choices": ["Yes", "No"],
        "label": "Include Testimony?"
    },
    # Conditional Testimony fields
    {
        "id": "testimony_before",
        "question": "How was the life of the participant before joining the support group?",
        "type": "text",
        "condition": lambda ans: ans.get("add_testimony") == "Yes",
        "label": "Participant Life Before"
    },
    {
        "id": "testimony_changes",
        "question": "What positive changes were noted in their life?",
        "type": "text",
        "condition": lambda ans: ans.get("add_testimony") == "Yes",
        "label": "Observed Changes"
    },
    {
        "id": "testimony_affirmations_status",
        "question": "Is the participant you are testifying about able to verbalize any positive affirmations?\n\n1. Yes\n2. No",
        "type": "choice",
        "choices": ["Yes", "No"],
        "condition": lambda ans: ans.get("add_testimony") == "Yes",
        "label": "Can Verbalize Affirmations?"
    },
    {
        "id": "testimony_affirmations",
        "question": "Please list the affirmations they verbalized:",
        "type": "text",
        "condition": lambda ans: ans.get("add_testimony") == "Yes" and ans.get("testimony_affirmations_status") == "Yes",
        "label": "Affirmations List"
    }
]

# Session memory store
# Session structure: { 'step_index': int, 'answers': dict, 'completed': bool }
sessions = {}

def serialize_step(step):
    """Helper to remove non-JSON serializable keys (like condition lambdas) from step dicts."""
    if step is None:
        return None
    return {k: v for k, v in step.items() if k != "condition"}

def get_next_step_index(answers, current_index):
    """Finds the index of the next step whose condition evaluates to True."""
    next_index = current_index + 1
    while next_index < len(FLOW):
        step = FLOW[next_index]
        if "condition" in step:
            try:
                # If condition evaluates to True, we found our next step
                if step["condition"](answers):
                    return next_index
            except Exception as e:
                print(f"Error evaluating condition for step {step['id']}: {e}")
        else:
            # No condition means it's a general step
            return next_index
        next_index += 1
    return len(FLOW)

def fill_docx_report(answers):
    """Populates the Word document template with responses and saves it."""
    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError(f"Template Word document not found at {TEMPLATE_PATH}")
    
    doc = Document(TEMPLATE_PATH)
    
    # 1. Fill Paragraph Placeholders (Header details)
    for p in doc.paragraphs:
        if "Group Facilitator" in p.text:
            p.text = f"Group Facilitator: {answers.get('facilitator', '')}        Cell #: {answers.get('cell_num', '')}"
        elif "Town/Village" in p.text:
            p.text = f"Town/Village: {answers.get('town_village', '')}        Location: {answers.get('location', '')}"
        elif "Meeting Day" in p.text:
            p.text = f"Meeting Day: {answers.get('meeting_day', '')}        Time: {answers.get('meeting_time', '')}"
            
    # 2. Fill Table 0 (Meeting logs)
    if len(doc.tables) > 0:
        t0 = doc.tables[0]
        
        # Month (Row 0, Cell 1)
        t0.rows[0].cells[1].text = answers.get('month', '')
        
        # Met Status Checkboxes (Row 1)
        # Yes -> Cell 3, No -> Cell 5
        if answers.get('met_status') == "Yes":
            t0.rows[1].cells[3].text = "  [✓] Yes"
            t0.rows[1].cells[5].text = "  [ ] No"
        else:
            t0.rows[1].cells[3].text = "  [ ] Yes"
            t0.rows[1].cells[5].text = "  [✓] No"
            
        # Attendees (Row 2)
        # Male -> Cell 3, Female -> Cell 5
        if answers.get('met_status') == "Yes":
            t0.rows[2].cells[3].text = f" {answers.get('attendees_male', '0')} Male"
            t0.rows[2].cells[5].text = f" {answers.get('attendees_female', '0')} Female"
        else:
            t0.rows[2].cells[3].text = " 0 Male"
            t0.rows[2].cells[5].text = " 0 Female"
            
        # Lessons Interesting (Row 4, Cell 0)
        t0.rows[4].cells[0].text = answers.get('lessons_interesting', 'N/A - Did not meet')
        
        # Challenges Encountered (Row 6, Cell 0)
        challenges_list = answers.get('challenges', [])
        challenges_text = ", ".join(challenges_list)
        if "Other" in challenges_list and answers.get('challenges_other'):
            challenges_text += f" (Other: {answers.get('challenges_other')})"
        t0.rows[6].cells[0].text = challenges_text if challenges_text else "None"
        
        # Challenges Resolved (Row 8, Cell 0)
        t0.rows[8].cells[0].text = answers.get('challenges_resolved', 'None')
        
        # Challenges Unresolved (Row 10, Cell 0)
        t0.rows[10].cells[0].text = answers.get('challenges_unresolved', 'None')
        
    # 3. Fill Table 1 (Testimony Details)
    if len(doc.tables) > 1:
        t1 = doc.tables[1]
        if answers.get('add_testimony') == "Yes":
            t1.rows[1].cells[0].text = answers.get('testimony_before', '')
            t1.rows[3].cells[0].text = answers.get('testimony_changes', '')
            
            affirmations_text = ""
            if answers.get('testimony_affirmations_status') == "Yes":
                affirmations_text = f"Yes. Affirmations: {answers.get('testimony_affirmations', '')}"
            else:
                affirmations_text = "No affirmations verbalized."
            t1.rows[5].cells[0].text = affirmations_text
        else:
            t1.rows[1].cells[0].text = "No participant testimony attached for this reporting cycle."
            t1.rows[3].cells[0].text = "N/A"
            t1.rows[5].cells[0].text = "N/A"
            
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH

@app.route("/")
def index():
    # We serve the frontend HTML using render_template_string
    # We'll load the HTML file from the local workspace
    try:
        with open("index.html", "r", encoding="utf-8") as f:
            html_content = f.read()
        return render_template_string(html_content)
    except Exception as e:
        return f"Error loading index.html: {str(e)}"

@app.route("/api/reset", methods=["POST"])
def reset_session():
    session_id = request.json.get("session_id", "default")
    sessions[session_id] = {
        "step_index": 0,
        "answers": {},
        "completed": False
    }
    
    # Return the first question
    first_step = FLOW[0]
    return jsonify({
        "success": True,
        "message": first_step["question"],
        "step": serialize_step(first_step),
        "state": sessions[session_id]
    })

@app.route("/api/message", methods=["POST"])
def process_message():
    data = request.json
    session_id = data.get("session_id", "default")
    user_message = data.get("message", "").strip()
    
    if session_id not in sessions:
        sessions[session_id] = {
            "step_index": 0,
            "answers": {},
            "completed": False
        }
        return jsonify({
            "message": FLOW[0]["question"],
            "step": serialize_step(FLOW[0]),
            "state": sessions[session_id]
        })
        
    session = sessions[session_id]
    
    if session["completed"]:
        return jsonify({
            "message": "Your report is already fully compiled and generated! Click 'Download Word Document' on the panel to save it.",
            "completed": True,
            "state": session
        })
        
    current_index = session["step_index"]
    current_step = FLOW[current_index]
    
    # Validate and process the response for the CURRENT question
    raw_val = user_message
    
    if current_step["type"] == "choice":
        # Check if they inputted a number or the exact text
        choices = current_step["choices"]
        selected_choice = None
        
        # Try matching number
        if raw_val.isdigit():
            idx = int(raw_val) - 1
            if 0 <= idx < len(choices):
                selected_choice = choices[idx]
        else:
            # Try fuzzy match
            for choice in choices:
                if choice.lower() == raw_val.lower():
                    selected_choice = choice
                    break
                    
        if not selected_choice:
            # Invalid selection
            options_str = "\n".join([f"{i+1}. {ch}" for i, ch in enumerate(choices)])
            return jsonify({
                "message": f"⚠️ Invalid option. Please reply with a number (1-{len(choices)}) or choose one of the options below:\n\n{options_str}",
                "step": serialize_step(current_step),
                "state": session
            })
        raw_val = selected_choice
        
    elif current_step["type"] == "integer":
        if not raw_val.isdigit():
            return jsonify({
                "message": "⚠️ Please reply with a valid whole number (e.g. 15):",
                "step": serialize_step(current_step),
                "state": session
            })
        raw_val = int(raw_val)
        
    elif current_step["type"] == "multi_choice":
        # Multi-select checkbox simulation
        choices = current_step["choices"]
        selected_choices = []
        
        if raw_val.lower() == "none" or raw_val.lower() == "7" or raw_val == "None":
            selected_choices = []
        else:
            # Process comma-separated values (e.g., "1, 3, 5")
            parts = [p.strip() for p in raw_val.split(",") if p.strip()]
            for p in parts:
                if p.isdigit():
                    idx = int(p) - 1
                    if 0 <= idx < len(choices):
                        selected_choices.append(choices[idx])
                else:
                    # Search text
                    for choice in choices:
                        if choice.lower() == p.lower():
                            selected_choices.append(choice)
                            break
                            
            if not selected_choices and raw_val != "":
                return jsonify({
                    "message": "⚠️ I couldn't understand that. Please select options by number (e.g., 1, 3, 5) or type 'None':",
                    "step": serialize_step(current_step),
                    "state": session
                })
        raw_val = selected_choices

    # Save the answer
    session["answers"][current_step["id"]] = raw_val
    
    # Move to the next applicable step
    next_index = get_next_step_index(session["answers"], current_index)
    session["step_index"] = next_index
    
    if next_index >= len(FLOW):
        # We are finished! Generate document.
        session["completed"] = True
        try:
            fill_docx_report(session["answers"])
            msg = "🎉 Excellent! I have collected all the information and filled out your **HFF Support Group Report Form** successfully.\n\nI have generated the official Word Document for you. You can download it directly from the Developer Dashboard on the right!"
        except Exception as e:
            msg = f"🎉 I have collected all information, but there was an error generating the Word Document: {str(e)}"
            
        return jsonify({
            "message": msg,
            "completed": True,
            "state": session
        })
        
    next_step = FLOW[next_index]
    return jsonify({
        "message": next_step["question"],
        "step": serialize_step(next_step),
        "state": session
    })

@app.route("/api/download")
def download_report():
    if os.path.exists(OUTPUT_PATH):
        return send_file(OUTPUT_PATH, as_attachment=True, download_name="HFF_Support_Group_Report_Filled.docx")
    else:
        return "Report file not generated yet.", 404

if __name__ == "__main__":
    app.run(port=5000, debug=True)
