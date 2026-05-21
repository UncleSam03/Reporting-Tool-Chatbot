import os
import json
import tempfile
import requests
from flask import Flask, request, jsonify, send_file, render_template_string
from docx import Document
from dotenv import load_dotenv
from typing import Optional, List
from pydantic import BaseModel, Field

# LangChain Imports
from langchain_core.prompts import ChatPromptTemplate
from langchain_google_genai import ChatGoogleGenerativeAI

# Load environment variables
load_dotenv()

app = Flask(__name__)

# Constants
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH = os.path.join(BASE_DIR, "HFF SUPPORT GROUP REPORTING TOOL.docx")
OUTPUT_PATH = os.path.join(tempfile.gettempdir(), "HFF_Support_Group_Report_Filled.docx")

# Schema for Structured Entity Extraction
class ReportForm(BaseModel):
    facilitator: Optional[str] = Field(None, description="Name of the Group Facilitator.")
    cell_num: Optional[str] = Field(None, description="Cell/Phone number of the facilitator.")
    town_village: Optional[str] = Field(None, description="Town or Village where the support group is based.")
    location: Optional[str] = Field(None, description="Specific Location or Venue of meetings.")
    meeting_day: Optional[str] = Field(None, description="Day of the week meetings are held (Monday, Tuesday, Wednesday, Thursday, Friday, Saturday, Sunday).")
    meeting_time: Optional[str] = Field(None, description="Time of meetings (e.g., 10:00 AM or 14:00).")
    month: Optional[str] = Field(None, description="Month of the report (e.g., January, February, etc.).")
    met_status: Optional[str] = Field(None, description="Were they able to meet this month? Answer 'Yes' or 'No'.")
    attendees_male: Optional[int] = Field(None, description="Number of Male participants. Only if met_status is 'Yes'.")
    attendees_female: Optional[int] = Field(None, description="Number of Female participants. Only if met_status is 'Yes'.")
    lessons_interesting: Optional[str] = Field(None, description="Description of how lessons were made interesting. Only if met_status is 'Yes'.")
    challenges: Optional[List[str]] = Field(None, description="List of challenges encountered. Allowed items: 'Attendance / Punctuality issues', 'Lack of learning materials / manuals', 'Venue availability or bad weather', 'Disinterest / low participation from members', 'Transportation constraints', 'Healthcare or physical health barriers', 'Other'.")
    challenges_other: Optional[str] = Field(None, description="Description of other challenges. Only if challenges contains 'Other'.")
    challenges_resolved: Optional[str] = Field(None, description="How challenges were resolved.")
    challenges_unresolved: Optional[str] = Field(None, description="Suggestions on how unresolved challenges can be solved.")
    add_testimony: Optional[str] = Field(None, description="Would they like to add a participant testimony? Answer 'Yes' or 'No'.")
    testimony_before: Optional[str] = Field(None, description="Participant's life before joining the group. Only if add_testimony is 'Yes'.")
    testimony_changes: Optional[str] = Field(None, description="Positive changes observed in their life. Only if add_testimony is 'Yes'.")
    testimony_affirmations_status: Optional[str] = Field(None, description="Are they able to verbalize positive affirmations? Answer 'Yes' or 'No'. Only if add_testimony is 'Yes'.")
    testimony_affirmations: Optional[str] = Field(None, description="List of positive affirmations they verbalized. Only if testimony_affirmations_status is 'Yes'.")

FLOW_METADATA = {
    "facilitator": {
        "id": "facilitator",
        "question": "What is your name (Group Facilitator)?",
        "type": "text",
        "label": "Facilitator Name"
    },
    "cell_num": {
        "id": "cell_num",
        "question": "Thank you. What is your Cell # (Phone number)?",
        "type": "text",
        "label": "Cell Number"
    },
    "town_village": {
        "id": "town_village",
        "question": "Got it. Which Town or Village is this support group located in?",
        "type": "text",
        "label": "Town/Village"
    },
    "location": {
        "id": "location",
        "question": "What is the specific Location or Venue of the meetings?",
        "type": "text",
        "label": "Specific Location"
    },
    "meeting_day": {
        "id": "meeting_day",
        "question": "On which day of the week does the group meet?",
        "type": "choice",
        "choices": ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"],
        "label": "Meeting Day"
    },
    "meeting_time": {
        "id": "meeting_time",
        "question": "At what time do meetings start? (e.g. 10:00 AM or 14:00)",
        "type": "text",
        "label": "Meeting Time"
    },
    "month": {
        "id": "month",
        "question": "Which Month is this report for?",
        "type": "choice",
        "choices": ["January", "February", "March", "April", "May", "June", "July", "August", "September", "October", "November", "December"],
        "label": "Reporting Month"
    },
    "met_status": {
        "id": "met_status",
        "question": "Were you able to meet this month?",
        "type": "choice",
        "choices": ["Yes", "No"],
        "label": "Able to Meet?"
    },
    "attendees_male": {
        "id": "attendees_male",
        "question": "How many Male participants turned up?",
        "type": "integer",
        "label": "Male Attendees"
    },
    "attendees_female": {
        "id": "attendees_female",
        "question": "How many Female participants turned up?",
        "type": "integer",
        "label": "Female Attendees"
    },
    "lessons_interesting": {
        "id": "lessons_interesting",
        "question": "How do you make the lessons interesting for the participants?",
        "type": "text",
        "label": "Lessons Interest"
    },
    "challenges": {
        "id": "challenges",
        "question": "What challenges have you encountered? (Select all that apply or type 'None'):",
        "type": "multi_choice",
        "choices": [
            "Attendance / Punctuality issues",
            "Lack of learning materials / manuals",
            "Venue availability or bad weather",
            "Disinterest / low participation from members",
            "Transportation constraints",
            "Healthcare or physical health barriers",
            "Other"
        ],
        "label": "Challenges Encountered"
    },
    "challenges_other": {
        "id": "challenges_other",
        "question": "You selected 'Other'. Please describe the other challenges you encountered:",
        "type": "text",
        "label": "Other Challenges Description"
    },
    "challenges_resolved": {
        "id": "challenges_resolved",
        "question": "How were the challenges resolved?",
        "type": "text",
        "label": "How Challenges Resolved"
    },
    "challenges_unresolved": {
        "id": "challenges_unresolved",
        "question": "If the challenges haven't been resolved, how do you think they can be resolved?",
        "type": "text",
        "label": "Unresolved Resolution Ideas"
    },
    "add_testimony": {
        "id": "add_testimony",
        "question": "Would you like to add a Support Group Participant Testimony to this report?",
        "type": "choice",
        "choices": ["Yes", "No"],
        "label": "Include Testimony?"
    },
    "testimony_before": {
        "id": "testimony_before",
        "question": "How was the life of the participant before joining the support group?",
        "type": "text",
        "label": "Participant Life Before"
    },
    "testimony_changes": {
        "id": "testimony_changes",
        "question": "What positive changes were noted in their life?",
        "type": "text",
        "label": "Observed Changes"
    },
    "testimony_affirmations_status": {
        "id": "testimony_affirmations_status",
        "question": "Is the participant you are testifying about able to verbalize any positive affirmations?",
        "type": "choice",
        "choices": ["Yes", "No"],
        "label": "Can Verbalize Affirmations?"
    },
    "testimony_affirmations": {
        "id": "testimony_affirmations",
        "question": "Please list the affirmations they verbalized:",
        "type": "text",
        "label": "Affirmations List"
    }
}

# Active conversational sessions
sessions = {}

def get_missing_fields(answers):
    missing = []
    
    # 1. General core details
    core = ["facilitator", "cell_num", "town_village", "location", "meeting_day", "meeting_time", "month", "met_status"]
    for field in core:
        if field not in answers or answers[field] is None or answers[field] == "":
            missing.append(field)
            
    # 2. Conditional meeting logs (only if met_status == 'Yes')
    if answers.get("met_status") == "Yes":
        meeting = ["attendees_male", "attendees_female", "lessons_interesting"]
        for field in meeting:
            if field not in answers or answers[field] is None or answers[field] == "":
                missing.append(field)
                
    # 3. Challenges details
    if "challenges" not in answers or answers["challenges"] is None:
        missing.append("challenges")
    else:
        challenges_list = answers.get("challenges", [])
        if isinstance(challenges_list, list) and "Other" in challenges_list:
            if "challenges_other" not in answers or answers["challenges_other"] is None or answers["challenges_other"] == "":
                missing.append("challenges_other")
                
    # 4. Challenges resolution
    resolutions = ["challenges_resolved", "challenges_unresolved", "add_testimony"]
    for field in resolutions:
        if field not in answers or answers[field] is None or answers[field] == "":
            missing.append(field)
            
    # 5. Conditional testimony fields (only if add_testimony == 'Yes')
    if answers.get("add_testimony") == "Yes":
        testimony = ["testimony_before", "testimony_changes", "testimony_affirmations_status"]
        for field in testimony:
            if field not in answers or answers[field] is None or answers[field] == "":
                missing.append(field)
                
        if answers.get("testimony_affirmations_status") == "Yes":
            if "testimony_affirmations" not in answers or answers["testimony_affirmations"] is None or answers["testimony_affirmations"] == "":
                missing.append("testimony_affirmations")
                
    return missing

def fill_docx_report(answers):
    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError(f"Template Word document not found at {TEMPLATE_PATH}")
    
    doc = Document(TEMPLATE_PATH)
    
    for p in doc.paragraphs:
        if "Group Facilitator" in p.text:
            p.text = f"Group Facilitator: {answers.get('facilitator', '')}        Cell #: {answers.get('cell_num', '')}"
        elif "Town/Village" in p.text:
            p.text = f"Town/Village: {answers.get('town_village', '')}        Location: {answers.get('location', '')}"
        elif "Meeting Day" in p.text:
            p.text = f"Meeting Day: {answers.get('meeting_day', '')}        Time: {answers.get('meeting_time', '')}"
            
    if len(doc.tables) > 0:
        t0 = doc.tables[0]
        t0.rows[0].cells[1].text = answers.get('month', '')
        
        if answers.get('met_status') == "Yes":
            t0.rows[1].cells[3].text = "  [✓] Yes"
            t0.rows[1].cells[5].text = "  [ ] No"
        else:
            t0.rows[1].cells[3].text = "  [ ] Yes"
            t0.rows[1].cells[5].text = "  [✓] No"
            
        if answers.get('met_status') == "Yes":
            t0.rows[2].cells[3].text = f" {answers.get('attendees_male', '0')} Male"
            t0.rows[2].cells[5].text = f" {answers.get('attendees_female', '0')} Female"
        else:
            t0.rows[2].cells[3].text = " 0 Male"
            t0.rows[2].cells[5].text = " 0 Female"
            
        t0.rows[4].cells[0].text = answers.get('lessons_interesting', 'N/A - Did not meet')
        
        challenges_list = answers.get('challenges', [])
        challenges_text = ", ".join(challenges_list) if isinstance(challenges_list, list) else str(challenges_list)
        if "Other" in challenges_list and answers.get('challenges_other'):
            challenges_text += f" (Other: {answers.get('challenges_other')})"
        t0.rows[6].cells[0].text = challenges_text if challenges_text else "None"
        t0.rows[8].cells[0].text = answers.get('challenges_resolved', 'None')
        t0.rows[10].cells[0].text = answers.get('challenges_unresolved', 'None')
        
    if len(doc.tables) > 1:
        t1 = doc.tables[1]
        if answers.get('add_testimony') == "Yes":
            t1.rows[1].cells[0].text = answers.get('testimony_before', '')
            t1.rows[3].cells[0].text = answers.get('testimony_changes', '')
            
            affirmations_text = ""
            if answers.get('testimony_affirmations_status') == "Yes":
                affirmations_text = f"Yes. Affirmations: {answers.get('testimony_affirmations', '')}"
            else:
                affirmations_text = "No affirmations verbalized."
            t1.rows[5].cells[0].text = affirmations_text
        else:
            t1.rows[1].cells[0].text = "No participant testimony attached for this reporting cycle."
            t1.rows[3].cells[0].text = "N/A"
            t1.rows[5].cells[0].text = "N/A"
            
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH

# LangChain AI Helpers
extraction_prompt = ChatPromptTemplate.from_messages([
    ("system", """You are an expert data extraction agent. 
Your goal is to extract support group report fields from the chat conversation history.

Here is the current state of extracted fields (answers) so far:
{current_answers_json}

Carefully read the conversation history, especially the most recent messages. Update or fill in the report fields based ONLY on what the user has explicitly stated.
If the user corrects an answer (e.g., 'No, my name is Sarah' or 'Actually we met on Thursday'), update that field.
If the user chooses a pill/choice option or clicks a WhatsApp button, extract that value exactly.
If the user explicitly states they had no challenges or types 'None', set challenges to [].
Do not invent details. Leave fields as null if they have not been provided yet.
""")
])

conversational_prompt_template = """You are the HFF Support Group Report Assistant, a friendly and empathetic AI chatbot helping support group facilitators compile their monthly reporting data.

We are collecting data in a WhatsApp chat interface. Your goal is to guide the user naturally through filling out their report.

Current completed fields:
{current_answers_json}

The next field we MUST collect is:
- ID: {next_field_id}
- Label: {next_field_label}
- Standard Question: {next_field_question}

INSTRUCTIONS:
1. Review the conversation history.
2. If the user asked a question, had a comment, or said something off-topic, address it in a highly professional, clear, and empathetic way first.
3. Then, ask for the value of the missing field. Make it sound natural and conversational (not a dry form).
4. Keep your message relatively concise and formatting clean (WhatsApp friendly). Use bold text (*like this*) for emphasis. Note: WhatsApp uses a single asterisk (*text*) for bolding, not double asterisks.
5. DO NOT mention technical IDs or internal parameters (like '{next_field_id}') to the user.
6. Do not ask for multiple fields at once. Just focus on collecting '{next_field_label}'.
"""

def run_extraction(messages_history, current_answers):
    try:
        llm = ChatGoogleGenerativeAI(model="gemini-1.5-flash", temperature=0)
        structured_llm = llm.with_structured_output(ReportForm)
        
        system_content = extraction_prompt.format(current_answers_json=json.dumps(current_answers, indent=2))
        
        langchain_messages = [("system", system_content)]
        for msg in messages_history[-12:]:
            langchain_messages.append((msg["role"], msg["content"]))
            
        extracted_data = structured_llm.invoke(langchain_messages)
        
        updates = extracted_data.model_dump(exclude_unset=True, exclude_none=True)
        for k, v in updates.items():
            if v is not None:
                current_answers[k] = v
                
        # Handle conditional cleanup
        if current_answers.get("met_status") == "No":
            for k in ["attendees_male", "attendees_female", "lessons_interesting"]:
                current_answers.pop(k, None)
        if isinstance(current_answers.get("challenges"), list) and "Other" not in current_answers["challenges"]:
            current_answers.pop("challenges_other", None)
        if current_answers.get("add_testimony") == "No":
            for k in ["testimony_before", "testimony_changes", "testimony_affirmations_status", "testimony_affirmations"]:
                current_answers.pop(k, None)
        if current_answers.get("testimony_affirmations_status") == "No":
            current_answers.pop("testimony_affirmations", None)
    except Exception as e:
        print(f"Extraction error: {e}")

def generate_bot_response(messages_history, current_answers, next_field_info):
    try:
        llm = ChatGoogleGenerativeAI(model="gemini-1.5-flash", temperature=0.5)
        
        system_content = conversational_prompt_template.format(
            current_answers_json=json.dumps(current_answers, indent=2),
            next_field_id=next_field_info["id"],
            next_field_label=next_field_info["label"],
            next_field_question=next_field_info["question"]
        )
        
        langchain_messages = [("system", system_content)]
        for msg in messages_history[-8:]:
            langchain_messages.append((msg["role"], msg["content"]))
            
        res = llm.invoke(langchain_messages)
        return res.content
    except Exception as e:
        print(f"Generation error: {e}")
        return next_field_info["question"]

# WhatsApp API Handlers
def send_whatsapp_message(to_phone, text):
    token = os.environ.get("WHATSAPP_TOKEN")
    phone_id = os.environ.get("WHATSAPP_PHONE_NUMBER_ID")
    if not token or not phone_id:
        print("[WARNING] WhatsApp credentials not set in environment.")
        return False
        
    url = f"https://graph.facebook.com/v18.0/{phone_id}/messages"
    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json"
    }
    payload = {
        "messaging_product": "whatsapp",
        "recipient_type": "individual",
        "to": to_phone,
        "type": "text",
        "text": {"body": text}
    }
    
    res = requests.post(url, json=payload, headers=headers)
    return res.status_code == 200

def send_whatsapp_buttons(to_phone, text, choices):
    token = os.environ.get("WHATSAPP_TOKEN")
    phone_id = os.environ.get("WHATSAPP_PHONE_NUMBER_ID")
    if not token or not phone_id:
        return send_whatsapp_message(to_phone, text)
        
    # WhatsApp button messages can take up to 3 choices
    if len(choices) > 3 or len(choices) == 0:
        choices_str = "\n".join([f"• {ch}" for ch in choices])
        return send_whatsapp_message(to_phone, f"{text}\n\n{choices_str}")
        
    url = f"https://graph.facebook.com/v18.0/{phone_id}/messages"
    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json"
    }
    
    buttons = []
    for idx, choice in enumerate(choices):
        buttons.append({
            "type": "reply",
            "reply": {
                "id": f"btn_{idx}",
                "title": choice[:20]  # Max 20 characters
            }
        })
        
    payload = {
        "messaging_product": "whatsapp",
        "recipient_type": "individual",
        "to": to_phone,
        "type": "interactive",
        "interactive": {
            "type": "button",
            "body": {"text": text},
            "action": {"buttons": buttons}
        }
    }
    
    res = requests.post(url, json=payload, headers=headers)
    return res.status_code == 200

def upload_and_send_document(to_phone, file_path):
    token = os.environ.get("WHATSAPP_TOKEN")
    phone_id = os.environ.get("WHATSAPP_PHONE_NUMBER_ID")
    if not token or not phone_id:
        print("[WARNING] WhatsApp credentials missing. Cannot upload media.")
        return False
        
    # 1. Upload media
    upload_url = f"https://graph.facebook.com/v18.0/{phone_id}/media"
    headers = {
        "Authorization": f"Bearer {token}"
    }
    files = {
        "file": ("HFF_Support_Group_Report.docx", open(file_path, "rb"), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
        "messaging_product": (None, "whatsapp")
    }
    
    res = requests.post(upload_url, headers=headers, files=files)
    if res.status_code != 200:
        print(f"Media upload failed: {res.text}")
        return False
        
    media_id = res.json().get("id")
    
    # 2. Send media document
    send_url = f"https://graph.facebook.com/v18.0/{phone_id}/messages"
    send_headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json"
    }
    payload = {
        "messaging_product": "whatsapp",
        "recipient_type": "individual",
        "to": to_phone,
        "type": "document",
        "document": {
            "id": media_id,
            "filename": "HFF_Support_Group_Report_Filled.docx"
        }
    }
    
    send_res = requests.post(send_url, json=payload, headers=send_headers)
    return send_res.status_code == 200


# --- LIVE WHATSAPP WEBHOOK ROUTES ---

@app.route("/webhook", methods=["GET"])
def whatsapp_verify():
    """Meta webhook verification route (GET challenge handshake)."""
    verify_token = os.environ.get("WHATSAPP_VERIFY_TOKEN")
    
    mode = request.args.get("hub.mode")
    token = request.args.get("hub.verify_token")
    challenge = request.args.get("hub.challenge")
    
    if mode and token:
        if mode == "subscribe" and token == verify_token:
            print("[OK] Webhook verification SUCCESS!")
            return challenge, 200
        else:
            print("[ERROR] Webhook verification FAILED: Token mismatch.")
            return "Forbidden", 403
            
    return "Verification API Active", 200

@app.route("/webhook", methods=["POST"])
def whatsapp_webhook():
    """Live WhatsApp inbound messages webhook (POST event receiver)."""
    data = request.json
    print("Received WhatsApp Webhook Payload:", json.dumps(data, indent=2))
    
    # Extract message details
    try:
        entry = data.get("entry", [])[0]
        changes = entry.get("changes", [])[0]
        value = changes.get("value", {})
        messages = value.get("messages", [])
        
        if not messages:
            return "No messages to process", 200
            
        msg = messages[0]
        sender_phone = msg.get("from")
        msg_type = msg.get("type")
        
        # Parse text body depending on message category (text or interactive selection)
        user_text = ""
        if msg_type == "text":
            user_text = msg.get("text", {}).get("body", "")
        elif msg_type == "interactive":
            interactive = msg.get("interactive", {})
            int_type = interactive.get("type")
            if int_type == "button_reply":
                user_text = interactive.get("button_reply", {}).get("title", "")
            elif int_type == "list_reply":
                user_text = interactive.get("list_reply", {}).get("title", "")
                
        if not sender_phone or not user_text:
            return "Ignored message type", 200
            
        # Initialize or fetch session using sender's phone number as session ID!
        if sender_phone not in sessions:
            sessions[sender_phone] = {
                "answers": {},
                "completed": False,
                "messages": []
            }
            
        session = sessions[sender_phone]
        
        if session.get("completed"):
            send_whatsapp_message(
                sender_phone, 
                "Your monthly report is already compiled and finalized! If you need to make a new one, type 'reset'."
            )
            return "Done", 200
            
        # Handle manual reset trigger
        if user_text.strip().lower() == "reset":
            session["answers"] = {}
            session["completed"] = False
            session["messages"] = []
            user_text = "Hi" # Restart conversational extraction
            
        # 1. Update session history log
        session["messages"].append({"role": "user", "content": user_text})
        
        # 2. Extract state updates using LangChain structured output
        run_extraction(session["messages"], session["answers"])
        
        # 3. Assess unfilled/missing fields
        missing_fields = get_missing_fields(session["answers"])
        
        if not missing_fields:
            # Report is ready!
            session["completed"] = True
            try:
                doc_path = fill_docx_report(session["answers"])
                
                # Save completed report to SQLite database
                import database
                database.save_report(sender_phone, session["answers"])
                
                success_msg = "🎉 *Excellent!* I have collected all the information and filled out your *HFF Support Group Report Form*.\n\nGenerating your official Word Document report now..."
                send_whatsapp_message(sender_phone, success_msg)
                
                # Upload and send the file directly to WhatsApp!
                upload_and_send_document(sender_phone, doc_path)
            except Exception as e:
                send_whatsapp_message(sender_phone, f"🎉 I collected all your data, but faced an issue sending the document file: {e}")
                
            return "Done", 200
            
        # 4. Target the next missing field
        next_field_id = missing_fields[0]
        next_meta = FLOW_METADATA[next_field_id]
        
        step_data = {
            "id": next_field_id,
            "question": next_meta["question"],
            "type": next_meta["type"],
            "label": next_meta["label"]
        }
        if "choices" in next_meta:
            step_data["choices"] = next_meta["choices"]
            
        # 5. Generate conversational bot reply
        bot_reply = generate_bot_response(session["messages"], session["answers"], step_data)
        session["messages"].append({"role": "assistant", "content": bot_reply})
        
        # 6. Dispatch back to user (using buttons where appropriate)
        choices = next_meta.get("choices", [])
        # Only use buttons if there are 2 or 3 choices (e.g. Yes/No)
        if len(choices) in [2, 3]:
            send_whatsapp_buttons(sender_phone, bot_reply, choices)
        else:
            send_whatsapp_message(sender_phone, bot_reply)
            
    except Exception as e:
        print(f"Error handling WhatsApp webhook: {e}")
        
    return "Received", 200


# --- SIMULATOR & LOCAL TESTING ROUTES ---
# These routes remain intact so the dashboard client in the browser keeps working perfectly!

@app.route("/")
def index():
    try:
        with open("index.html", "r", encoding="utf-8") as f:
            html_content = f.read()
        html_content = html_content.replace(
            "Rule-Based WhatsApp Conversational Flow Prototype", 
            "LangChain-Powered WhatsApp & Simulator Engine"
        )
        return render_template_string(html_content)
    except Exception as e:
        return f"Error loading index.html: {str(e)}"

@app.route("/api/reset", methods=["POST"])
def reset_session():
    session_id = request.json.get("session_id", "default")
    
    if not os.environ.get("GEMINI_API_KEY"):
        return jsonify({
            "success": True,
            "message": "⚠️ **GEMINI_API_KEY Not Found**\n\nPlease set your `GEMINI_API_KEY` in a `.env` file in the project folder to start the LangChain AI agent.",
            "step": None,
            "state": {"step_index": 0, "answers": {}, "completed": False}
        })

    sessions[session_id] = {
        "answers": {},
        "completed": False,
        "messages": []
    }
    
    first_field_id = "facilitator"
    first_meta = FLOW_METADATA[first_field_id]
    initial_prompt = "Hello! Welcome to the HFF Support Group Reporting Chatbot. Let's start compiling your monthly report.\n\nWhat is your name (Group Facilitator)?"
    
    sessions[session_id]["messages"].append({"role": "assistant", "content": initial_prompt})
    
    step_data = {
        "id": first_field_id,
        "question": initial_prompt,
        "type": first_meta["type"],
        "label": first_meta["label"]
    }
    if "choices" in first_meta:
        step_data["choices"] = first_meta["choices"]
        
    return jsonify({
        "success": True,
        "message": initial_prompt,
        "step": step_data,
        "state": {"step_index": 0, "answers": {}, "completed": False}
    })

@app.route("/api/message", methods=["POST"])
def process_message():
    data = request.json
    session_id = data.get("session_id", "default")
    user_message = data.get("message", "").strip()
    
    if not os.environ.get("GEMINI_API_KEY"):
        return jsonify({
            "message": "⚠️ **GEMINI_API_KEY Missing** in .env",
            "completed": False,
            "state": {"step_index": 0, "answers": {}, "completed": False}
        })
        
    if session_id not in sessions:
        return reset_session()
        
    session = sessions[session_id]
    
    if session.get("completed"):
        return jsonify({
            "message": "Your report is already fully compiled and generated! Click 'Download Word Document' on the panel to save it.",
            "completed": True,
            "state": {
                "step_index": len(FLOW_METADATA),
                "answers": session["answers"],
                "completed": True
            }
        })
        
    session["messages"].append({"role": "user", "content": user_message})
    run_extraction(session["messages"], session["answers"])
    missing_fields = get_missing_fields(session["answers"])
    
    total_fields = len(FLOW_METADATA)
    filled_count = total_fields - len(missing_fields)
    
    if not missing_fields:
        session["completed"] = True
        try:
            fill_docx_report(session["answers"])
            
            # Save completed report to SQLite database
            import database
            database.save_report(session_id, session["answers"])
            
            msg = "🎉 Excellent! I have collected all the information and filled out your **HFF Support Group Report Form** successfully.\n\nI have generated the official Word Document for you. You can download it directly from the Developer Dashboard on the right!"
        except Exception as e:
            msg = f"🎉 I have collected all information, but there was an error generating the Word Document: {str(e)}"
            
        session["messages"].append({"role": "assistant", "content": msg})
        
        return jsonify({
            "message": msg,
            "completed": True,
            "state": {
                "step_index": total_fields,
                "answers": session["answers"],
                "completed": True
            }
        })
        
    next_field_id = missing_fields[0]
    next_meta = FLOW_METADATA[next_field_id]
    
    step_data = {
        "id": next_field_id,
        "question": next_meta["question"],
        "type": next_meta["type"],
        "label": next_meta["label"]
    }
    if "choices" in next_meta:
        step_data["choices"] = next_meta["choices"]
        
    bot_reply = generate_bot_response(session["messages"], session["answers"], step_data)
    session["messages"].append({"role": "assistant", "content": bot_reply})
    
    return jsonify({
        "message": bot_reply,
        "step": step_data,
        "state": {
            "step_index": filled_count,
            "answers": session["answers"],
            "completed": False
        }
    })

@app.route("/api/reports", methods=["GET"])
def get_reports():
    import database
    reports = database.get_all_reports()
    return jsonify(reports)

@app.route("/api/download")
def download_report():
    if os.path.exists(OUTPUT_PATH):
        return send_file(OUTPUT_PATH, as_attachment=True, download_name="HFF_Support_Group_Report_Filled.docx")
    else:
        return "Report file not generated yet.", 404

if __name__ == "__main__":
    # Standard Flask listener on port 5000 (accessible locally and globally via ngrok)
    app.run(port=5000, debug=True)
